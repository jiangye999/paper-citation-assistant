#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
参考文献反插助手 - PyQt5 桌面版
严格参考UI版逻辑实现
"""

import sys
import os
from pathlib import Path
from datetime import datetime

sys.path.insert(0, str(Path(__file__).parent / "src"))

from PyQt5.QtWidgets import (
    QApplication,
    QMainWindow,
    QWidget,
    QVBoxLayout,
    QHBoxLayout,
    QLabel,
    QPushButton,
    QTabWidget,
    QTextEdit,
    QFileDialog,
    QProgressBar,
    QMessageBox,
    QGroupBox,
    QFormLayout,
    QLineEdit,
    QComboBox,
    QSlider,
    QScrollArea,
    QFrame,
    QSplitter,
    QListWidget,
    QListWidgetItem,
    QCheckBox,
    QSpinBox,
    QPlainTextEdit,
    QStatusBar,
    QTableWidget,
    QTableWidgetItem,
    QHeaderView,
    QToolButton,
    QMenu,
    QAction,
    QSystemTrayIcon,
    QStyle,
)
from PyQt5.QtCore import Qt, QThread, pyqtSignal, QSettings, QSize
from PyQt5.QtGui import QFont, QIcon, QPalette, QColor

from src.literature.db_manager import LiteratureDatabaseManager
from src.draft.analyzer import DraftAnalyzer
from src.citation.ai_matcher import AICitationMatcher, AIAPIManager
from src.citation.format_learner import ReferenceFormatLearner


class ImportWorker(QThread):
    progress = pyqtSignal(int, str)
    finished = pyqtSignal(object, int, list)
    error = pyqtSignal(str)

    def __init__(self, files, db_path):
        super().__init__()
        self.files = files
        self.db_path = db_path

    def run(self):
        try:
            db_manager = LiteratureDatabaseManager(self.db_path)
            total_count = 0
            all_errors = []

            for idx, file_path in enumerate(self.files):
                self.progress.emit(
                    int((idx + 1) / len(self.files) * 100),
                    f"正在导入: {Path(file_path).name}",
                )

                count, errors = db_manager.import_from_wos_txt(file_path)
                total_count += count
                all_errors.extend(errors)

            self.finished.emit(db_manager, total_count, all_errors)
        except Exception as e:
            self.error.emit(str(e))


class AnalysisWorker(QThread):
    finished = pyqtSignal(object)
    error = pyqtSignal(str)

    def __init__(self, file_path):
        super().__init__()
        self.file_path = file_path

    def run(self):
        try:
            analyzer = DraftAnalyzer()
            analysis = analyzer.analyze_draft(self.file_path)
            self.finished.emit(analysis)
        except Exception as e:
            self.error.emit(str(e))


class MatchWorker(QThread):
    progress = pyqtSignal(int, str)
    finished = pyqtSignal(list)
    error = pyqtSignal(str)

    def __init__(self, db_manager, sentences, config):
        super().__init__()
        self.db_manager = db_manager
        self.sentences = sentences
        self.config = config

    def run(self):
        try:
            api_manager = AIAPIManager(
                api_key=self.config["api_key"],
                base_url=self.config.get("api_base_url", "https://api.deepseek.com/v1"),
                model=self.config.get("model", "deepseek-chat"),
                provider=self.config.get("api_provider", "deepseek"),
            )

            matcher = AICitationMatcher(
                db_manager=self.db_manager,
                api_manager=api_manager,
                citation_style=self.config["citation_style"],
                max_citations=self.config["max_citations"],
                min_relevance=self.config["min_relevance"],
                batch_size=5,
                top_k_semantic=self.config.get("top_k_semantic", 50),
                weight_recency=self.config.get("weight_recency", 50),
                weight_citation=self.config.get("weight_citation", 50),
                use_hybrid_search=self.config.get("use_hybrid_search", True),
            )

            def progress_callback(current, total):
                self.progress.emit(
                    int(current / total * 100), f"正在AI匹配: 句子 {current}/{total}"
                )

            results = matcher.batch_match(
                sentences=self.sentences,
                year_range=self.config.get("year_range", 10),
                progress_callback=progress_callback,
            )

            self.finished.emit(results)
        except Exception as e:
            self.error.emit(str(e))


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("参考文献反插助手")
        self.setMinimumSize(1400, 900)

        # 初始化状态
        self.db_manager = None
        self.draft_analysis = None
        self.citation_results = None
        self.imported_files = []

        # 默认配置
        self.config = {
            "api_provider": "deepseek",
            "api_key": "",
            "api_base_url": "https://api.deepseek.com/v1",
            "model": "deepseek-chat",
            "citation_style": "author-year",
            "max_citations": 2,
            "min_relevance": 0.6,
            "top_k_semantic": 50,
            "weight_recency": 50,
            "weight_citation": 50,
            "use_hybrid_search": True,
            "year_range": 10,
            "reference_format": None,
        }

        self.settings = QSettings("PaperCitation", "参考文献反插助手")
        self.load_settings()
        self.init_ui()
        self.apply_styles()

    def load_settings(self):
        self.config["api_provider"] = self.settings.value("api_provider", "deepseek")
        self.config["api_key"] = self.settings.value("api_key", "")
        self.config["model"] = self.settings.value("model", "deepseek-chat")
        self.config["citation_style"] = self.settings.value(
            "citation_style", "author-year"
        )
        self.config["max_citations"] = int(self.settings.value("max_citations", 2))
        self.config["min_relevance"] = float(self.settings.value("min_relevance", 0.6))

    def save_settings(self):
        self.settings.setValue("api_provider", self.config["api_provider"])
        self.settings.setValue("api_key", self.config["api_key"])
        self.settings.setValue("model", self.config["model"])
        self.settings.setValue("citation_style", self.config["citation_style"])
        self.settings.setValue("max_citations", self.config["max_citations"])
        self.settings.setValue("min_relevance", self.config["min_relevance"])

    def init_ui(self):
        # 中央部件
        central = QWidget()
        self.setCentralWidget(central)
        main_layout = QHBoxLayout(central)
        main_layout.setSpacing(0)
        main_layout.setContentsMargins(0, 0, 0, 0)

        # 左侧主区域
        left_widget = QWidget()
        left_layout = QVBoxLayout(left_widget)
        left_layout.setContentsMargins(20, 20, 20, 20)
        left_layout.setSpacing(15)

        # 标题
        title_layout = QHBoxLayout()
        title_icon = QLabel("📖")
        title_icon.setFont(QFont("Segoe UI Emoji", 32))
        title_layout.addWidget(title_icon)

        title_text = QLabel("参考文献反插助手")
        title_text.setFont(QFont("Microsoft YaHei", 24, QFont.Bold))
        title_layout.addWidget(title_text)
        title_layout.addStretch()

        subtitle = QLabel("基于AI的学术论文引用自动插入工具")
        subtitle.setFont(QFont("Microsoft YaHei", 11))
        subtitle.setStyleSheet("color: #666;")

        left_layout.addLayout(title_layout)
        left_layout.addWidget(subtitle)
        left_layout.addSpacing(20)

        # 标签页
        self.tabs = QTabWidget()
        self.tabs.setDocumentMode(True)

        # 创建四个标签页
        self.tab_import = LiteratureImportTab(self)
        self.tab_draft = DraftUploadTab(self)
        self.tab_match = CitationMatchingTab(self)
        self.tab_results = ResultsReviewTab(self)

        self.tabs.addTab(self.tab_import, "📚 导入文献库")
        self.tabs.addTab(self.tab_draft, "📝 上传草稿")
        self.tabs.addTab(self.tab_match, "🔍 AI匹配引用")
        self.tabs.addTab(self.tab_results, "📊 查看与导出")

        left_layout.addWidget(self.tabs)

        # 右侧边栏
        self.sidebar = Sidebar(self)
        self.sidebar.setFixedWidth(350)

        # 分割器
        splitter = QSplitter(Qt.Horizontal)
        splitter.addWidget(left_widget)
        splitter.addWidget(self.sidebar)
        splitter.setStretchFactor(0, 1)
        splitter.setStretchFactor(1, 0)
        splitter.setHandleWidth(1)

        main_layout.addWidget(splitter)

        # 状态栏
        self.statusbar = QStatusBar()
        self.setStatusBar(self.statusbar)
        self.statusbar.showMessage("就绪")

        # 更新数据库状态显示
        self.update_db_status()

    def apply_styles(self):
        self.setStyleSheet("""
            QMainWindow {
                background-color: #f8f9fa;
            }
            QTabWidget::pane {
                border: 1px solid #dee2e6;
                background-color: white;
                border-radius: 8px;
            }
            QTabBar::tab {
                background-color: #e9ecef;
                padding: 12px 24px;
                margin-right: 4px;
                border-top-left-radius: 8px;
                border-top-right-radius: 8px;
                font-size: 13px;
                font-weight: 500;
            }
            QTabBar::tab:selected {
                background-color: white;
                border-bottom: 3px solid #28a745;
            }
            QTabBar::tab:hover:!selected {
                background-color: #dee2e6;
            }
            QPushButton {
                background-color: #28a745;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 6px;
                font-size: 13px;
                font-weight: 500;
            }
            QPushButton:hover {
                background-color: #218838;
            }
            QPushButton:pressed {
                background-color: #1e7e34;
            }
            QPushButton:disabled {
                background-color: #6c757d;
            }
            QPushButton#secondary {
                background-color: #6c757d;
            }
            QPushButton#secondary:hover {
                background-color: #5a6268;
            }
            QGroupBox {
                border: 1px solid #dee2e6;
                border-radius: 8px;
                margin-top: 12px;
                padding-top: 12px;
                padding: 15px;
                font-weight: 600;
                background-color: white;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 15px;
                padding: 0 8px;
                color: #495057;
            }
            QLineEdit, QTextEdit, QPlainTextEdit, QComboBox {
                border: 1px solid #ced4da;
                border-radius: 6px;
                padding: 8px;
                background-color: white;
            }
            QLineEdit:focus, QTextEdit:focus, QPlainTextEdit:focus {
                border: 2px solid #28a745;
            }
            QProgressBar {
                border: 1px solid #dee2e6;
                border-radius: 4px;
                text-align: center;
                height: 20px;
            }
            QProgressBar::chunk {
                background-color: #28a745;
                border-radius: 4px;
            }
            QSlider::groove:horizontal {
                height: 6px;
                background: #dee2e6;
                border-radius: 3px;
            }
            QSlider::handle:horizontal {
                width: 18px;
                height: 18px;
                background: #28a745;
                border-radius: 9px;
                margin: -6px 0;
            }
            QListWidget {
                border: 1px solid #dee2e6;
                border-radius: 6px;
                padding: 5px;
                background-color: white;
            }
            QListWidget::item {
                padding: 8px;
                border-radius: 4px;
            }
            QListWidget::item:selected {
                background-color: #d4edda;
                color: #155724;
            }
            QStatusBar {
                background-color: #f8f9fa;
                border-top: 1px solid #dee2e6;
            }
            QLabel#info {
                color: #0c5460;
                background-color: #d1ecf1;
                padding: 12px;
                border-radius: 6px;
                border-left: 4px solid #17a2b8;
            }
            QLabel#success {
                color: #155724;
                background-color: #d4edda;
                padding: 12px;
                border-radius: 6px;
                border-left: 4px solid #28a745;
            }
            QLabel#warning {
                color: #856404;
                background-color: #fff3cd;
                padding: 12px;
                border-radius: 6px;
                border-left: 4px solid #ffc107;
            }
        """)

    def update_db_status(self):
        if self.db_manager:
            stats = self.db_manager.get_statistics()
            self.statusbar.showMessage(
                f"文献库: {stats['total_papers']} 篇 | 年份: {stats.get('earliest_year', '-')} - {stats.get('latest_year', '-')}"
            )
        else:
            self.statusbar.showMessage("就绪 - 请先导入文献库")


class Sidebar(QWidget):
    def __init__(self, main_window):
        super().__init__()
        self.main_window = main_window
        self.setStyleSheet("background-color: #f8f9fa;")
        self.init_ui()

    def init_ui(self):
        # 主布局
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(10, 10, 10, 10)
        main_layout.setSpacing(10)

        # 标题
        title_label = QLabel("⚙️ 配置面板")
        title_label.setFont(QFont("Microsoft YaHei", 14, QFont.Bold))
        title_label.setStyleSheet("color: #333; padding-bottom: 10px;")
        main_layout.addWidget(title_label)

        # 创建滚动区域
        self.scroll = QScrollArea()
        self.scroll.setWidgetResizable(True)
        self.scroll.setFrameShape(QScrollArea.NoFrame)
        self.scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self.scroll.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        self.scroll.setStyleSheet(
            "QScrollArea { border: none; background: transparent; }"
        )

        # 创建内容容器
        container = QWidget()
        container.setStyleSheet("background-color: transparent;")
        layout = QVBoxLayout(container)
        layout.setSpacing(12)
        layout.setContentsMargins(5, 5, 5, 5)

        # API配置
        api_group = QGroupBox("🔑 API配置")
        api_group.setStyleSheet(
            "QGroupBox { font-weight: bold; margin-top: 10px; } QGroupBox::title { subcontrol-origin: margin; left: 10px; padding: 0 5px; }"
        )
        api_layout = QFormLayout()
        api_layout.setSpacing(8)
        api_layout.setContentsMargins(10, 15, 10, 10)

        self.api_provider = QComboBox()
        self.api_provider.addItems(["deepseek", "openai", "anthropic"])
        self.api_provider.setCurrentText(self.main_window.config["api_provider"])
        self.api_provider.currentTextChanged.connect(self.on_provider_changed)
        api_layout.addRow("提供商:", self.api_provider)

        self.api_key = QLineEdit()
        self.api_key.setEchoMode(QLineEdit.Password)
        self.api_key.setText(self.main_window.config["api_key"])
        self.api_key.setPlaceholderText("sk-...")
        api_layout.addRow("API密钥:", self.api_key)

        self.model = QComboBox()
        self.update_model_list()
        self.model.setCurrentText(self.main_window.config["model"])
        api_layout.addRow("模型:", self.model)

        # API状态 - 增大padding和行高确保文字显示完整
        self.api_status = QLabel("⚠️ 请输入API密钥")
        self.api_status.setStyleSheet(
            "color: #856404; padding: 12px; background: #fff3cd; border-radius: 6px; font-size: 13px; min-height: 20px;"
        )
        self.api_status.setWordWrap(True)
        api_layout.addRow(self.api_status)

        api_group.setLayout(api_layout)
        layout.addWidget(api_group)

        # 引用设置
        citation_group = QGroupBox("📚 引用设置")
        citation_layout = QFormLayout()
        citation_layout.setSpacing(10)

        self.citation_style = QComboBox()
        self.citation_style.addItems(["author-year", "numbered"])
        self.citation_style.setCurrentText(self.main_window.config["citation_style"])
        citation_layout.addRow("引用风格:", self.citation_style)

        self.max_citations = QSpinBox()
        self.max_citations.setRange(1, 5)
        self.max_citations.setValue(self.main_window.config["max_citations"])
        citation_layout.addRow("每句最大引用:", self.max_citations)

        self.min_relevance = QSlider(Qt.Horizontal)
        self.min_relevance.setRange(0, 100)
        self.min_relevance.setValue(int(self.main_window.config["min_relevance"] * 100))
        self.min_relevance_label = QLabel(f"{self.min_relevance.value()}%")
        self.min_relevance.valueChanged.connect(self.update_relevance_label)
        relevance_layout = QHBoxLayout()
        relevance_layout.addWidget(self.min_relevance)
        relevance_layout.addWidget(self.min_relevance_label)
        citation_layout.addRow("最低相关性:", relevance_layout)

        citation_group.setLayout(citation_layout)
        layout.addWidget(citation_group)

        # 检索引擎设置
        search_group = QGroupBox("🔍 检索引擎")
        search_layout = QVBoxLayout()

        self.use_hybrid = QCheckBox("启用混合检索")
        self.use_hybrid.setChecked(self.main_window.config["use_hybrid_search"])
        self.use_hybrid.setToolTip("启用AI增强的混合检索（需要模型文件）")
        search_layout.addWidget(self.use_hybrid)

        self.hybrid_info = QLabel(
            "✅ 查询扩展 → 多路召回 → Cross-encoder重排 → MMR多样"
        )
        self.hybrid_info.setStyleSheet("color: #28a745; font-size: 11px;")
        search_layout.addWidget(self.hybrid_info)

        self.top_k = QSpinBox()
        self.top_k.setRange(10, 100)
        self.top_k.setValue(self.main_window.config["top_k_semantic"])
        self.top_k.setSingleStep(10)
        search_layout.addWidget(QLabel("语义筛选保留数量:"))
        search_layout.addWidget(self.top_k)

        search_group.setLayout(search_layout)
        layout.addWidget(search_group)

        # 权重设置
        weight_group = QGroupBox("⚖️ 文献筛选策略")
        weight_layout = QVBoxLayout()

        weight_layout.addWidget(QLabel("新颖度 vs 引用权重:"))

        weight_slider_layout = QHBoxLayout()
        self.weight_recency = QSlider(Qt.Horizontal)
        self.weight_recency.setRange(0, 100)
        self.weight_recency.setValue(self.main_window.config["weight_recency"])
        self.weight_recency.valueChanged.connect(self.update_weight_labels)
        weight_slider_layout.addWidget(self.weight_recency)

        self.weight_label = QLabel(
            f"新颖度 {self.weight_recency.value()}% | 引用 {100 - self.weight_recency.value()}%"
        )
        weight_slider_layout.addWidget(self.weight_label)
        weight_layout.addLayout(weight_slider_layout)

        # 预设按钮
        preset_layout = QHBoxLayout()
        btn_balanced = QPushButton("⚖️ 均衡")
        btn_balanced.clicked.connect(lambda: self.weight_recency.setValue(50))
        btn_new = QPushButton("🆕 追新")
        btn_new.clicked.connect(lambda: self.weight_recency.setValue(80))
        preset_layout.addWidget(btn_balanced)
        preset_layout.addWidget(btn_new)
        weight_layout.addLayout(preset_layout)

        weight_group.setLayout(weight_layout)
        layout.addWidget(weight_group)

        # 数据库状态
        self.db_status_group = QGroupBox("📊 数据库状态")
        db_layout = QVBoxLayout()

        self.db_status_label = QLabel("⚠️ 未导入文献")
        self.db_status_label.setStyleSheet(
            "padding: 10px; background: #fff3cd; border-radius: 6px; color: #856404;"
        )
        db_layout.addWidget(self.db_status_label)

        self.db_status_group.setLayout(db_layout)
        layout.addWidget(self.db_status_group)

        # 保存按钮
        save_btn = QPushButton("💾 保存所有设置")
        save_btn.setStyleSheet("padding: 12px; font-size: 14px;")
        save_btn.clicked.connect(self.save_config)
        layout.addWidget(save_btn)

        layout.addStretch()

        # 设置滚动区域的内容
        self.scroll.setWidget(container)

        # 将滚动区域添加到主布局
        main_layout.addWidget(self.scroll)

        # 检查API状态
        self.check_api_status()
        self.api_key.textChanged.connect(self.on_api_key_changed)

    def on_api_key_changed(self):
        """API密钥改变时更新状态并同步到主窗口"""
        self.check_api_status()
        # 实时同步到主窗口配置
        self.main_window.config["api_key"] = self.api_key.text()
        # 通知匹配标签页更新状态
        if hasattr(self.main_window, "tab_match"):
            self.main_window.tab_match.update_api_status()

    def update_model_list(self):
        provider = self.api_provider.currentText()
        self.model.clear()
        if provider == "openai":
            self.model.addItems(["gpt-4o", "gpt-4o-mini", "gpt-4-turbo"])
        elif provider == "anthropic":
            self.model.addItems(["claude-3-5-sonnet-20241022"])
        else:
            self.model.addItems(["deepseek-chat", "deepseek-reasoner"])

    def on_provider_changed(self):
        self.update_model_list()
        self.check_api_status()

    def check_api_status(self):
        if self.api_key.text():
            self.api_status.setText("✅ API已配置")
            self.api_status.setStyleSheet(
                "color: #155724; padding: 8px; background: #d4edda; border-radius: 4px;"
            )
        else:
            self.api_status.setText("⚠️ 请输入API密钥")
            self.api_status.setStyleSheet(
                "color: #856404; padding: 8px; background: #fff3cd; border-radius: 4px;"
            )

    def update_relevance_label(self):
        self.min_relevance_label.setText(f"{self.min_relevance.value()}%")

    def update_weight_labels(self):
        self.weight_label.setText(
            f"新颖度 {self.weight_recency.value()}% | 引用 {100 - self.weight_recency.value()}%"
        )

    def update_db_status(self, stats):
        if stats:
            text = f"✅ 已加载文献库\n"
            text += f"文献数量: {stats['total_papers']}\n"
            text += f"年份范围: {stats.get('earliest_year', '-')} - {stats.get('latest_year', '-')}"
            self.db_status_label.setText(text)
            self.db_status_label.setStyleSheet(
                "padding: 10px; background: #d4edda; border-radius: 6px; color: #155724;"
            )

    def save_config(self):
        self.main_window.config["api_provider"] = self.api_provider.currentText()
        self.main_window.config["api_key"] = self.api_key.text()
        self.main_window.config["model"] = self.model.currentText()
        self.main_window.config["citation_style"] = self.citation_style.currentText()
        self.main_window.config["max_citations"] = self.max_citations.value()
        self.main_window.config["min_relevance"] = self.min_relevance.value() / 100.0
        self.main_window.config["top_k_semantic"] = self.top_k.value()
        self.main_window.config["weight_recency"] = self.weight_recency.value()
        self.main_window.config["weight_citation"] = 100 - self.weight_recency.value()
        self.main_window.config["use_hybrid_search"] = self.use_hybrid.isChecked()

        self.main_window.save_settings()
        QMessageBox.information(self, "成功", "✅ 所有设置已保存")


class LiteratureImportTab(QWidget):
    def __init__(self, main_window):
        super().__init__()
        self.main_window = main_window
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout(self)
        layout.setSpacing(20)
        layout.setContentsMargins(20, 20, 20, 20)

        # 标题
        title = QLabel("📚 导入文献库")
        title.setFont(QFont("Microsoft YaHei", 18, QFont.Bold))
        layout.addWidget(title)

        # 操作指南
        guide = QLabel(
            "<b>📋 操作指南</b><br>"
            "<b>从Web of Science导入：</b><br>"
            "1. 在Web of Science中搜索文献<br>"
            "2. 选择要导出的文献 → 3. 点击 <b>Export</b> → <b>Plain Text File</b><br>"
            "4. 选择 <b>Full Record</b> 格式 → 5. 下载 .txt 文件<br>"
            "6. 在下方上传文件"
        )
        guide.setStyleSheet(
            "background: #e7f3ff; padding: 15px; border-radius: 8px; border-left: 4px solid #2196F3;"
        )
        guide.setWordWrap(True)
        layout.addWidget(guide)

        # 文件上传区域
        upload_group = QGroupBox("上传WOS导出文件")
        upload_layout = QVBoxLayout()

        self.file_list = QListWidget()
        self.file_list.setMinimumHeight(150)
        upload_layout.addWidget(QLabel("已选择的文件:"))
        upload_layout.addWidget(self.file_list)

        btn_layout = QHBoxLayout()

        self.btn_add = QPushButton("📁 添加文件")
        self.btn_add.clicked.connect(self.add_files)
        btn_layout.addWidget(self.btn_add)

        self.btn_clear = QPushButton("🗑️ 清空")
        self.btn_clear.clicked.connect(self.clear_files)
        btn_layout.addWidget(self.btn_clear)

        btn_layout.addStretch()

        self.btn_import = QPushButton("🚀 开始导入")
        self.btn_import.setStyleSheet("padding: 12px 30px; font-size: 14px;")
        self.btn_import.clicked.connect(self.start_import)
        btn_layout.addWidget(self.btn_import)

        upload_layout.addLayout(btn_layout)

        # 进度区域
        self.progress_widget = QWidget()
        progress_layout = QVBoxLayout(self.progress_widget)

        self.progress_label = QLabel("")
        progress_layout.addWidget(self.progress_label)

        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        progress_layout.addWidget(self.progress_bar)

        self.progress_widget.setVisible(False)
        upload_layout.addWidget(self.progress_widget)

        upload_group.setLayout(upload_layout)
        layout.addWidget(upload_group)

        # 结果显示区域
        self.result_widget = QWidget()
        result_layout = QVBoxLayout(self.result_widget)

        self.result_label = QLabel("")
        self.result_label.setStyleSheet(
            "padding: 15px; background: #d4edda; border-radius: 8px; color: #155724;"
        )
        result_layout.addWidget(self.result_label)

        # 统计卡片
        stats_layout = QHBoxLayout()

        self.stat_papers = self.create_stat_card("总文献数", "0")
        self.stat_years = self.create_stat_card("年份范围", "-")
        self.stat_journals = self.create_stat_card("期刊种类", "0")

        stats_layout.addWidget(self.stat_papers)
        stats_layout.addWidget(self.stat_years)
        stats_layout.addWidget(self.stat_journals)

        result_layout.addLayout(stats_layout)

        self.result_widget.setVisible(False)
        layout.addWidget(self.result_widget)

        layout.addStretch()

    def create_stat_card(self, title, value):
        card = QGroupBox()
        card_layout = QVBoxLayout()

        title_label = QLabel(title)
        title_label.setStyleSheet("color: #6c757d; font-size: 12px;")
        title_label.setAlignment(Qt.AlignCenter)

        value_label = QLabel(value)
        value_label.setFont(QFont("Microsoft YaHei", 24, QFont.Bold))
        value_label.setStyleSheet("color: #28a745;")
        value_label.setAlignment(Qt.AlignCenter)
        value_label.setObjectName(f"stat_{title}")

        card_layout.addWidget(title_label)
        card_layout.addWidget(value_label)
        card.setLayout(card_layout)

        return card

    def add_files(self):
        files, _ = QFileDialog.getOpenFileNames(
            self, "选择WOS导出文件", "", "Text Files (*.txt)"
        )
        for f in files:
            self.file_list.addItem(f)
        self.update_import_button()

    def clear_files(self):
        self.file_list.clear()
        self.update_import_button()

    def update_import_button(self):
        self.btn_import.setEnabled(self.file_list.count() > 0)

    def start_import(self):
        if self.file_list.count() == 0:
            return

        files = [self.file_list.item(i).text() for i in range(self.file_list.count())]
        db_path = str(Path(__file__).parent / "data" / "literature.db")

        # 确保目录存在
        os.makedirs(Path(db_path).parent, exist_ok=True)

        # 显示进度
        self.progress_widget.setVisible(True)
        self.progress_bar.setVisible(True)
        self.btn_import.setEnabled(False)
        self.btn_add.setEnabled(False)
        self.btn_clear.setEnabled(False)

        # 创建工作线程
        self.worker = ImportWorker(files, db_path)
        self.worker.progress.connect(self.on_progress)
        self.worker.finished.connect(self.on_finished)
        self.worker.error.connect(self.on_error)
        self.worker.start()

    def on_progress(self, value, message):
        self.progress_bar.setValue(value)
        self.progress_label.setText(message)

    def on_finished(self, db_manager, total_count, errors):
        self.progress_bar.setVisible(False)
        self.btn_import.setEnabled(True)
        self.btn_add.setEnabled(True)
        self.btn_clear.setEnabled(True)

        # 保存数据库管理器
        self.main_window.db_manager = db_manager

        # 显示结果
        self.result_widget.setVisible(True)
        self.result_label.setText(f"✅ 成功导入 {total_count} 篇论文！")

        # 更新统计
        stats = db_manager.get_statistics()
        self.stat_papers.findChild(QLabel, "stat_总文献数").setText(
            str(stats["total_papers"])
        )

        years = list(stats.get("year_distribution", {}).keys())
        if years:
            year_range = f"{min(years)}-{max(years)}"
        else:
            year_range = "-"
        self.stat_years.findChild(QLabel, "stat_年份范围").setText(year_range)

        self.stat_journals.findChild(QLabel, "stat_期刊种类").setText(
            str(len(stats.get("top_journals", [])))
        )

        # 更新侧边栏状态
        self.main_window.sidebar.update_db_status(stats)
        self.main_window.update_db_status()

        # 显示错误
        if errors:
            error_msg = "\n".join(errors[:10])
            if len(errors) > 10:
                error_msg += f"\n... 还有 {len(errors) - 10} 个错误"
            QMessageBox.warning(
                self, "导入警告", f"导入完成，但有 {len(errors)} 个错误:\n{error_msg}"
            )

    def on_error(self, error):
        self.progress_bar.setVisible(False)
        self.btn_import.setEnabled(True)
        self.btn_add.setEnabled(True)
        self.btn_clear.setEnabled(True)
        QMessageBox.critical(self, "导入失败", f"导入过程中出错:\n{error}")


class DraftUploadTab(QWidget):
    def __init__(self, main_window):
        super().__init__()
        self.main_window = main_window
        self.file_path = None
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout(self)
        layout.setSpacing(20)
        layout.setContentsMargins(20, 20, 20, 20)

        # 标题
        title = QLabel("📝 上传草稿")
        title.setFont(QFont("Microsoft YaHei", 18, QFont.Bold))
        layout.addWidget(title)

        # 检查文献库状态
        self.status_label = QLabel("")
        self.status_label.setWordWrap(True)
        layout.addWidget(self.status_label)

        # 文件上传区域
        upload_group = QGroupBox("上传Word文档")
        upload_layout = QVBoxLayout()

        self.file_info = QLabel("未选择文件")
        self.file_info.setStyleSheet(
            "padding: 20px; background: #f8f9fa; border-radius: 8px; border: 2px dashed #dee2e6;"
        )
        self.file_info.setAlignment(Qt.AlignCenter)
        upload_layout.addWidget(self.file_info)

        btn_layout = QHBoxLayout()
        btn_layout.addStretch()

        self.btn_upload = QPushButton("📄 选择Word文档")
        self.btn_upload.clicked.connect(self.upload_file)
        btn_layout.addWidget(self.btn_upload)

        self.btn_analyze = QPushButton("🔬 分析文档")
        self.btn_analyze.setStyleSheet("padding: 12px 30px; font-size: 14px;")
        self.btn_analyze.clicked.connect(self.analyze_document)
        self.btn_analyze.setEnabled(False)
        btn_layout.addWidget(self.btn_analyze)

        btn_layout.addStretch()
        upload_layout.addLayout(btn_layout)

        # 进度
        self.progress_label = QLabel("")
        upload_layout.addWidget(self.progress_label)

        upload_group.setLayout(upload_layout)
        layout.addWidget(upload_group)

        # 分析结果
        self.result_group = QGroupBox("分析结果")
        result_layout = QVBoxLayout()

        # 统计卡片
        stats_layout = QHBoxLayout()

        self.stat_total = self.create_stat_card("总句子数", "0")
        self.stat_need = self.create_stat_card("需引用句子", "0")
        self.stat_para = self.create_stat_card("段落数", "0")

        stats_layout.addWidget(self.stat_total)
        stats_layout.addWidget(self.stat_need)
        stats_layout.addWidget(self.stat_para)

        result_layout.addLayout(stats_layout)

        # 句子预览
        result_layout.addWidget(QLabel("句子预览（前5句）:"))
        self.preview_text = QTextEdit()
        self.preview_text.setReadOnly(True)
        self.preview_text.setMaximumHeight(150)
        result_layout.addWidget(self.preview_text)

        self.result_group.setLayout(result_layout)
        self.result_group.setVisible(False)
        layout.addWidget(self.result_group)

        layout.addStretch()

        self.update_status()

    def create_stat_card(self, title, value):
        card = QGroupBox()
        card_layout = QVBoxLayout()

        title_label = QLabel(title)
        title_label.setStyleSheet("color: #6c757d; font-size: 12px;")
        title_label.setAlignment(Qt.AlignCenter)

        value_label = QLabel(value)
        value_label.setFont(QFont("Microsoft YaHei", 24, QFont.Bold))
        value_label.setStyleSheet("color: #28a745;")
        value_label.setAlignment(Qt.AlignCenter)

        card_layout.addWidget(title_label)
        card_layout.addWidget(value_label)
        card.setLayout(card_layout)

        return card

    def update_status(self):
        if self.main_window.db_manager:
            stats = self.main_window.db_manager.get_statistics()
            self.status_label.setText(
                f"✅ 已加载文献库: {stats['total_papers']} 篇论文"
            )
            self.status_label.setStyleSheet(
                "color: #155724; background: #d4edda; padding: 15px; border-radius: 8px; border-left: 4px solid #28a745;"
            )
            self.btn_upload.setEnabled(True)
        else:
            self.status_label.setText(
                '⚠️ 请先在左侧"导入文献库"中导入文献\n\n'
                "快速开始：\n"
                "1. 切换到 📚 导入文献库 标签\n"
                "2. 上传Web of Science导出的.txt文件\n"
                "3. 等待导入完成"
            )
            self.status_label.setStyleSheet(
                "color: #856404; background: #fff3cd; padding: 15px; border-radius: 8px; border-left: 4px solid #ffc107;"
            )
            self.btn_upload.setEnabled(False)

    def showEvent(self, event):
        """标签页显示时更新状态"""
        super().showEvent(event)
        self.update_status()

    def upload_file(self):
        file, _ = QFileDialog.getOpenFileName(
            self, "选择Word文档", "", "Word Documents (*.docx)"
        )
        if file:
            self.file_path = file
            file_size = Path(file).stat().st_size / 1024
            self.file_info.setText(f"📄 {Path(file).name}\n({file_size:.1f} KB)")
            self.file_info.setStyleSheet(
                "padding: 20px; background: #e7f3ff; border-radius: 8px; border: 2px solid #2196F3;"
            )
            self.btn_analyze.setEnabled(True)

    def analyze_document(self):
        if not self.file_path:
            return

        self.btn_analyze.setEnabled(False)
        self.btn_upload.setEnabled(False)
        self.progress_label.setText("正在分析文档结构...")

        # 创建工作线程
        self.worker = AnalysisWorker(self.file_path)
        self.worker.finished.connect(self.on_analysis_finished)
        self.worker.error.connect(self.on_analysis_error)
        self.worker.start()

    def on_analysis_finished(self, analysis):
        self.main_window.draft_analysis = analysis

        self.progress_label.setText("")
        self.btn_analyze.setEnabled(True)
        self.btn_upload.setEnabled(True)

        # 显示结果
        self.result_group.setVisible(True)

        # 更新统计
        total = len(analysis.sentences)
        need = len([s for s in analysis.sentences if not s.has_citation])
        paras = len(analysis.paragraphs)

        self.stat_total.findChildren(QLabel)[1].setText(str(total))
        self.stat_need.findChildren(QLabel)[1].setText(str(need))
        self.stat_para.findChildren(QLabel)[1].setText(str(paras))

        # 显示预览
        preview_text = ""
        for i, sent in enumerate(analysis.sentences[:5]):
            preview_text += f"<b>句子 {i + 1}:</b> {sent.text[:100]}...\n"
            if sent.keywords:
                preview_text += f"<span style='color: #6c757d;'>关键词: {', '.join(sent.keywords)}</span>\n\n"
        self.preview_text.setHtml(preview_text)

        QMessageBox.information(
            self, "分析完成", f"✅ 分析完成！\n总句子数: {total}\n需引用句子: {need}"
        )

    def on_analysis_error(self, error):
        self.progress_label.setText("")
        self.btn_analyze.setEnabled(True)
        self.btn_upload.setEnabled(True)
        QMessageBox.critical(self, "分析失败", f"分析过程中出错:\n{error}")


class CitationMatchingTab(QWidget):
    def __init__(self, main_window):
        super().__init__()
        self.main_window = main_window
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout(self)
        layout.setSpacing(20)
        layout.setContentsMargins(20, 20, 20, 20)

        # 标题
        title = QLabel("🔍 AI匹配引用")
        title.setFont(QFont("Microsoft YaHei", 18, QFont.Bold))
        layout.addWidget(title)

        # 前置条件检查
        self.check_label = QLabel("")
        self.check_label.setWordWrap(True)
        layout.addWidget(self.check_label)

        # 匹配选项
        options_group = QGroupBox("匹配选项")
        options_layout = QHBoxLayout()

        self.chk_skip_existing = QCheckBox("跳过已有引用的句子")
        self.chk_skip_existing.setChecked(True)
        self.chk_skip_existing.setToolTip("不处理已经包含引用的句子")
        options_layout.addWidget(self.chk_skip_existing)

        options_layout.addWidget(QLabel("文献年份范围:"))
        self.year_range = QSpinBox()
        self.year_range.setRange(5, 30)
        self.year_range.setValue(10)
        self.year_range.setToolTip("只搜索最近N年的文献")
        options_layout.addWidget(self.year_range)

        self.chk_prioritize_recent = QCheckBox("优先推荐新文献")
        self.chk_prioritize_recent.setChecked(True)
        self.chk_prioritize_recent.setToolTip("优先匹配近5年的文献")
        options_layout.addWidget(self.chk_prioritize_recent)

        options_layout.addStretch()
        options_group.setLayout(options_layout)
        layout.addWidget(options_group)

        # 开始匹配按钮
        btn_layout = QHBoxLayout()
        btn_layout.addStretch()

        self.btn_match = QPushButton("🚀 开始AI匹配引用")
        self.btn_match.setStyleSheet("padding: 15px 40px; font-size: 16px;")
        self.btn_match.clicked.connect(self.start_matching)
        btn_layout.addWidget(self.btn_match)

        btn_layout.addStretch()
        layout.addLayout(btn_layout)

        # 进度区域
        self.progress_widget = QWidget()
        progress_layout = QVBoxLayout(self.progress_widget)

        self.progress_info = QLabel("🤖 正在使用AI进行语义匹配，这可能需要一些时间...")
        progress_layout.addWidget(self.progress_info)

        self.progress_bar = QProgressBar()
        self.progress_bar.setTextVisible(True)
        progress_layout.addWidget(self.progress_bar)

        self.progress_status = QLabel("")
        progress_layout.addWidget(self.progress_status)

        self.progress_widget.setVisible(False)
        layout.addWidget(self.progress_widget)

        layout.addStretch()

    def showEvent(self, event):
        super().showEvent(event)
        self.check_prerequisites()

    def update_api_status(self):
        """当API配置改变时更新状态显示"""
        self.check_prerequisites()

    def check_prerequisites(self):
        issues = []

        if not self.main_window.db_manager:
            issues.append("⚠️ 请先导入文献库")

        if not self.main_window.draft_analysis:
            issues.append("⚠️ 请先上传并分析草稿")

        if not self.main_window.config.get("api_key"):
            issues.append("⚠️ 请在侧边栏输入API密钥")

        if issues:
            self.check_label.setText("<br>".join(issues))
            self.check_label.setStyleSheet(
                "color: #856404; background: #fff3cd; padding: 15px; border-radius: 8px; border-left: 4px solid #ffc107;"
            )
            self.btn_match.setEnabled(False)
        else:
            analysis = self.main_window.draft_analysis
            sentences = [s for s in analysis.sentences if not s.has_citation]
            self.check_label.setText(
                f"✅ 准备就绪<br>文献库: {self.main_window.db_manager.get_statistics()['total_papers']} 篇 | "
                f"需匹配句子: {len(sentences)} 句"
            )
            self.check_label.setStyleSheet(
                "color: #155724; background: #d4edda; padding: 15px; border-radius: 8px; border-left: 4px solid #28a745;"
            )
            self.btn_match.setEnabled(True)

    def start_matching(self):
        self.check_prerequisites()
        if not self.btn_match.isEnabled():
            return

        analysis = self.main_window.draft_analysis
        sentences = analysis.sentences

        if self.chk_skip_existing.isChecked():
            sentences = [s for s in sentences if not s.has_citation]

        if not sentences:
            QMessageBox.warning(self, "无匹配内容", "没有需要匹配的句子")
            return

        # 更新配置
        self.main_window.config["year_range"] = self.year_range.value()

        # 显示进度
        self.progress_widget.setVisible(True)
        self.btn_match.setEnabled(False)

        # 创建工作线程
        self.worker = MatchWorker(
            self.main_window.db_manager, sentences, self.main_window.config
        )
        self.worker.progress.connect(self.on_progress)
        self.worker.finished.connect(self.on_finished)
        self.worker.error.connect(self.on_error)
        self.worker.start()

    def on_progress(self, value, status):
        self.progress_bar.setValue(value)
        self.progress_status.setText(status)

    def on_finished(self, results):
        self.main_window.citation_results = results

        self.progress_widget.setVisible(False)
        self.btn_match.setEnabled(True)

        QMessageBox.information(
            self, "匹配完成", f"✅ AI匹配完成！\n共处理 {len(results)} 个句子"
        )

        # 自动切换到结果页
        self.main_window.tabs.setCurrentIndex(3)
        self.main_window.tab_results.refresh_results()

    def on_error(self, error):
        self.progress_widget.setVisible(False)
        self.btn_match.setEnabled(True)
        QMessageBox.critical(self, "匹配失败", f"匹配过程中出错:\n{error}")


class ResultsReviewTab(QWidget):
    def __init__(self, main_window):
        super().__init__()
        self.main_window = main_window
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout(self)
        layout.setSpacing(20)
        layout.setContentsMargins(20, 20, 20, 20)

        # 标题
        title = QLabel("📊 查看与导出")
        title.setFont(QFont("Microsoft YaHei", 18, QFont.Bold))
        layout.addWidget(title)

        # 导出选项
        export_group = QGroupBox("导出选项")
        export_layout = QHBoxLayout()

        export_layout.addWidget(QLabel("输出格式:"))
        self.export_format = QComboBox()
        self.export_format.addItems(
            ["Word文档 (.docx)", "Markdown (.md)", "纯文本 (.txt)"]
        )
        export_layout.addWidget(self.export_format)

        export_layout.addStretch()

        self.btn_refresh = QPushButton("🔄 刷新结果")
        self.btn_refresh.clicked.connect(self.refresh_results)
        export_layout.addWidget(self.btn_refresh)

        self.btn_export = QPushButton("💾 导出文档")
        self.btn_export.setStyleSheet("padding: 10px 30px;")
        self.btn_export.clicked.connect(self.export_document)
        export_layout.addWidget(self.btn_export)

        export_group.setLayout(export_layout)
        layout.addWidget(export_group)

        # 结果表格
        self.table = QTableWidget()
        self.table.setColumnCount(4)
        self.table.setHorizontalHeaderLabels(["序号", "句子", "引用数", "操作"])
        self.table.horizontalHeader().setStretchLastSection(True)
        self.table.horizontalHeader().setSectionResizeMode(1, QHeaderView.Stretch)
        self.table.setSelectionBehavior(QTableWidget.SelectRows)
        self.table.setAlternatingRowColors(True)
        layout.addWidget(self.table)

        # 提示
        self.hint_label = QLabel("请先完成AI匹配引用")
        self.hint_label.setAlignment(Qt.AlignCenter)
        self.hint_label.setStyleSheet("color: #6c757d; padding: 50px;")
        layout.addWidget(self.hint_label)

        self.update_visibility()

    def update_visibility(self):
        has_results = bool(self.main_window.citation_results)
        self.table.setVisible(has_results)
        self.hint_label.setVisible(not has_results)
        self.btn_export.setEnabled(has_results)

    def refresh_results(self):
        results = self.main_window.citation_results

        self.update_visibility()

        if not results:
            return

        self.table.setRowCount(len(results))

        for i, result in enumerate(results):
            # 序号
            self.table.setItem(i, 0, QTableWidgetItem(str(i + 1)))

            # 句子
            text = (
                result.sentence.text[:100] + "..."
                if len(result.sentence.text) > 100
                else result.sentence.text
            )
            self.table.setItem(i, 1, QTableWidgetItem(text))

            # 引用数
            citation_count = len(result.citations) if result.citations else 0
            item = QTableWidgetItem(str(citation_count))
            if citation_count > 0:
                item.setBackground(QColor(212, 237, 218))  # 绿色背景
            self.table.setItem(i, 2, item)

            # 操作按钮
            btn = QPushButton("查看详情")
            btn.clicked.connect(lambda checked, r=result: self.show_detail(r))
            self.table.setCellWidget(i, 3, btn)

    def show_detail(self, result):
        dialog = QDialog(self)
        dialog.setWindowTitle("引用详情")
        dialog.setMinimumSize(600, 400)

        layout = QVBoxLayout(dialog)

        # 句子
        sent_label = QLabel(f"<b>句子:</b><br>{result.sentence.text}")
        sent_label.setWordWrap(True)
        layout.addWidget(sent_label)

        # 引用列表
        if result.citations:
            layout.addWidget(QLabel(f"<b>推荐引用 ({len(result.citations)}篇):</b>"))

            for j, citation in enumerate(result.citations[:5], 1):
                ref_text = (
                    f"{j}. {citation.title[:80]}..."
                    if len(citation.title) > 80
                    else f"{j}. {citation.title}"
                )
                ref_label = QLabel(f"  {ref_text}")
                ref_label.setStyleSheet("color: #495057; margin-left: 20px;")
                layout.addWidget(ref_label)
        else:
            layout.addWidget(
                QLabel("<span style='color: #dc3545;'>未找到匹配的引用</span>")
            )

        # 关闭按钮
        btn = QPushButton("关闭")
        btn.clicked.connect(dialog.accept)
        layout.addWidget(btn)

        dialog.exec_()

    def export_document(self):
        if not self.main_window.citation_results:
            QMessageBox.warning(self, "无数据", "没有可导出的结果")
            return

        file_filter = {
            0: "Word Documents (*.docx)",
            1: "Markdown Files (*.md)",
            2: "Text Files (*.txt)",
        }

        file_path, _ = QFileDialog.getSaveFileName(
            self, "保存文档", "", file_filter[self.export_format.currentIndex()]
        )

        if not file_path:
            return

        try:
            results = self.main_window.citation_results
            config = self.main_window.config

            # 重建段落结构
            paragraph_map = {}
            for result in results:
                para_idx = result.sentence.paragraph_index
                if para_idx not in paragraph_map:
                    paragraph_map[para_idx] = []
                paragraph_map[para_idx].append(result)

            # 获取引用风格
            citation_style = config.get("citation_style", "author-year")
            ref_numbering = (
                "numbered" if citation_style == "numbered" else "author_year"
            )

            # 准备导出内容
            full_text = ""
            bibliography = ""

            # 创建 matcher 用于插入引用
            from src.citation.ai_matcher import AICitationMatcher, AIAPIManager

            api_manager = None
            if config.get("api_key"):
                api_manager = AIAPIManager(
                    api_key=config["api_key"],
                    base_url=config.get("api_base_url", "https://api.deepseek.com/v1"),
                    model=config.get("model", "deepseek-chat"),
                    provider=config.get("api_provider", "deepseek"),
                )

            matcher = None
            if self.main_window.db_manager and api_manager:
                matcher = AICitationMatcher(
                    db_manager=self.main_window.db_manager,
                    api_manager=api_manager,
                    citation_style=citation_style,
                )

            # 按段落组织文本
            for para_idx in sorted(paragraph_map.keys()):
                para_sentences = paragraph_map[para_idx]
                para_text_parts = []

                for result in para_sentences:
                    if (
                        result.citations
                        and not result.sentence.has_citation
                        and matcher
                    ):
                        new_text = matcher.insert_citations_into_text(
                            result.sentence, result.citations
                        )
                        para_text_parts.append(new_text)
                    else:
                        para_text_parts.append(result.sentence.text)

                full_text += " ".join(para_text_parts) + "\n\n"

            # 生成参考文献
            if matcher:
                used_papers = {}
                for swc in results:
                    for citation in swc.citations:
                        paper_id = citation.paper.id
                        if paper_id not in used_papers:
                            used_papers[paper_id] = citation.paper

                if used_papers:
                    sorted_papers = sorted(
                        used_papers.values(),
                        key=lambda p: (
                            p.authors.split(",")[0].strip().split()[-1]
                            if p.authors
                            else ""
                        ).lower(),
                    )

                    # 格式化参考文献
                    formatted_refs = []
                    for paper in sorted_papers:
                        authors = paper.authors.replace(";", ", ")
                        ref = f"{authors} ({paper.year}). {paper.title}. {paper.journal}, {paper.volume}({paper.issue}), {paper.pages}."
                        formatted_refs.append(ref)

                    # 根据序号格式添加
                    bibliography = "# References\n\n"
                    for i, ref in enumerate(formatted_refs, 1):
                        if ref_numbering == "numbered":
                            bibliography += f"[{i}] {ref}\n\n"
                        else:
                            bibliography += f"{ref}\n\n"
                else:
                    bibliography = "# References\n\n暂无引用文献"

            full_text += "\n" + bibliography.strip()

            # 根据格式导出
            format_idx = self.export_format.currentIndex()

            if format_idx == 2:  # 纯文本
                with open(file_path, "w", encoding="utf-8") as f:
                    f.write(full_text)
                QMessageBox.information(
                    self, "导出成功", f"✅ 文本文件已保存:\n{file_path}"
                )

            elif format_idx == 1:  # Markdown
                with open(file_path, "w", encoding="utf-8") as f:
                    f.write(full_text)
                QMessageBox.information(
                    self, "导出成功", f"✅ Markdown文件已保存:\n{file_path}"
                )

            else:  # Word文档
                from docx import Document
                from docx.shared import Pt
                from docx.oxml.ns import qn

                doc = Document()

                def set_times_new_roman(run):
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    run.font.size = Pt(12)

                # 添加内容（保持段落结构）
                for para_idx in sorted(paragraph_map.keys()):
                    para_sentences = paragraph_map[para_idx]
                    para_text_parts = []

                    for result in para_sentences:
                        if (
                            result.citations
                            and not result.sentence.has_citation
                            and matcher
                        ):
                            new_text = matcher.insert_citations_into_text(
                                result.sentence, result.citations
                            )
                            para_text_parts.append(new_text)
                        else:
                            para_text_parts.append(result.sentence.text)

                    paragraph_text = " ".join(para_text_parts)
                    p = doc.add_paragraph(paragraph_text)

                    for run in p.runs:
                        set_times_new_roman(run)

                # 添加参考文献
                if used_papers:
                    doc.add_heading("References", level=1)

                    for i, ref in enumerate(formatted_refs, 1):
                        if ref_numbering == "numbered":
                            p = doc.add_paragraph(f"[{i}] {ref}")
                        else:
                            p = doc.add_paragraph(ref)

                        for run in p.runs:
                            set_times_new_roman(run)

                doc.save(file_path)
                QMessageBox.information(
                    self, "导出成功", f"✅ Word文档已保存:\n{file_path}"
                )

        except Exception as e:
            QMessageBox.critical(self, "导出失败", f"导出过程中出错:\n{str(e)}")
            import traceback

            traceback.print_exc()


def main():
    app = QApplication(sys.argv)
    app.setApplicationName("参考文献反插助手")
    app.setOrganizationName("PaperCitation")

    # 设置应用样式
    app.setStyle("Fusion")

    # 设置字体
    font = QFont("Microsoft YaHei", 10)
    app.setFont(font)

    window = MainWindow()
    window.show()

    sys.exit(app.exec_())


if __name__ == "__main__":
    main()
