import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import nltk
from docx import Document


@dataclass
class Sentence:
    """句子数据类"""

    text: str
    index: int  # 句子在文档中的序号
    paragraph_index: int  # 所在段落序号
    keywords: List[str] = field(default_factory=list)
    has_citation: bool = False  # 是否已有引用
    citation_text: str = ""  # 已有的引用文本


@dataclass
class DraftAnalysisResult:
    """草稿分析结果"""

    sentences: List[Sentence] = field(default_factory=list)
    full_text: str = ""
    title: str = ""
    paragraphs: List[str] = field(default_factory=list)


class DraftAnalyzer:
    """草稿分析器"""

    def __init__(self):
        """初始化分析器"""
        self._ensure_nltk_data()

    def _ensure_nltk_data(self):
        """确保NLTK数据已下载"""
        try:
            nltk.data.find("tokenizers/punkt")
        except LookupError:
            try:
                nltk.download("punkt", quiet=True)
            except:
                pass

    def analyze_draft(self, file_path: str) -> DraftAnalysisResult:
        """
        分析草稿文档

        Args:
            file_path: Word文档路径

        Returns:
            草稿分析结果
        """
        result = DraftAnalysisResult()

        # 读取Word文档
        doc = Document(file_path)

        # 提取标题（通常是第一个段落）
        if doc.paragraphs:
            result.title = doc.paragraphs[0].text.strip()

        # 提取所有段落
        full_paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                full_paragraphs.append(text)

        result.paragraphs = full_paragraphs
        result.full_text = "\n\n".join(full_paragraphs)

        # 分割句子
        sentence_idx = 0
        for para_idx, paragraph in enumerate(full_paragraphs):
            sentences = self._split_sentences(paragraph)

            for sent_text in sentences:
                # 检查是否已有引用
                has_citation, citation_text = self._detect_citation(sent_text)

                # 提取关键词
                keywords = self._extract_keywords(sent_text)

                sentence = Sentence(
                    text=sent_text,
                    index=sentence_idx,
                    paragraph_index=para_idx,
                    keywords=keywords,
                    has_citation=has_citation,
                    citation_text=citation_text,
                )

                result.sentences.append(sentence)
                sentence_idx += 1

        return result

    def _split_sentences(self, text: str) -> List[str]:
        """
        智能分割句子

        考虑：
        - 缩写（如 e.g., i.e., Fig., et al.）
        - 小数点
        - 引号内的句子
        """
        if not text:
            return []

        try:
            # 尝试使用NLTK
            sentences = nltk.sent_tokenize(text)
        except:
            # 回退到简单分割
            sentences = self._simple_sentence_split(text)

        # 清理句子
        cleaned = []
        for sent in sentences:
            sent = sent.strip()
            if len(sent) >= 10:  # 过滤太短的句子
                cleaned.append(sent)

        return cleaned

    def _simple_sentence_split(self, text: str) -> List[str]:
        """简单的句子分割（当NLTK不可用时）"""
        # 保护常见的缩写和小数
        protected_patterns = [
            (r"\b(e\.g\.)\b", "___EG___"),
            (r"\b(i\.e\.)\b", "___IE___"),
            (r"\b(Fig\.\s*\d+)\b", "___FIG___"),
            (r"\b(Table\.?\s*\d+)\b", "___TABLE___"),
            (r"\b(et\s+al\.)\b", "___ETAL___"),
            (r"\b(vs\.)\b", "___VS___"),
            (r"\b(dr\.)\b", "___DR___"),
            (r"\b(mr\.)\b", "___MR___"),
            (r"\b(mrs\.)\b", "___MRS___"),
            (r"\b(st\.)\b", "___ST___"),
            (r"(\d+\.\d+)", "___DECIMAL___"),
        ]

        protected_text = text
        replacements = {}

        for pattern, placeholder in protected_patterns:
            matches = list(re.finditer(pattern, protected_text, re.IGNORECASE))
            for i, match in enumerate(matches):
                unique_placeholder = f"{placeholder}_{i}___"
                replacements[unique_placeholder] = match.group(1)
                protected_text = (
                    protected_text[: match.start()]
                    + unique_placeholder
                    + protected_text[match.end() :]
                )

        # 分割句子
        sentences = re.split(r"(?<=[.!?])\s+(?=[A-Z])", protected_text)

        # 恢复保护的内容
        restored_sentences = []
        for sent in sentences:
            for placeholder, original in replacements.items():
                sent = sent.replace(placeholder, original)
            restored_sentences.append(sent)

        return restored_sentences

    def _detect_citation(self, text: str) -> Tuple[bool, str]:
        """
        检测句子中是否已有引用

        Returns:
            (是否有引用, 引用文本)
        """
        # 作者-年份格式: (Author, Year) 或 (Author et al., Year)
        author_year_pattern = r"\([A-Z][a-z]+(?:\s+et\s+al\.?)?,\s*\d{4}[a-z]?\)"

        # 编号格式: [1] 或 [1-3] 或 [1, 2, 3]
        numbered_pattern = r"\[\d+(?:\s*[,-]\s*\d+)*\]"

        # 上标格式 (简化检测)
        superscript_pattern = r"\^\{\d+\}"

        for pattern in [author_year_pattern, numbered_pattern, superscript_pattern]:
            match = re.search(pattern, text)
            if match:
                return True, match.group(0)

        return False, ""

    def _extract_keywords(self, text: str, max_keywords: int = 5) -> List[str]:
        """
        从句子中提取关键词

        简单实现：提取长度适中的名词性词汇
        """
        # 清理文本
        text = re.sub(r"[^\w\s]", " ", text)

        # 分词
        words = text.split()

        # 过滤停用词和短词
        stopwords = {
            "the",
            "a",
            "an",
            "and",
            "or",
            "but",
            "in",
            "on",
            "at",
            "to",
            "for",
            "of",
            "with",
            "by",
            "from",
            "as",
            "is",
            "was",
            "are",
            "were",
            "been",
            "be",
            "have",
            "has",
            "had",
            "do",
            "does",
            "did",
            "will",
            "would",
            "could",
            "should",
            "may",
            "might",
            "can",
            "this",
            "that",
            "these",
            "those",
            "we",
            "our",
            "us",
            "it",
            "its",
            "they",
            "their",
            "them",
            "study",
            "research",
            "paper",
            "work",
            "results",
            "shown",
            "showed",
            "found",
            "observed",
            "indicated",
            "suggested",
            "demonstrated",
        }

        # 提取候选关键词（长度4-20的单词）
        candidates = []
        for word in words:
            word_lower = word.lower()
            if (
                4 <= len(word) <= 20
                and word_lower not in stopwords
                and word[0].isalpha()
            ):
                candidates.append(word_lower)

        # 统计词频
        word_freq = {}
        for word in candidates:
            word_freq[word] = word_freq.get(word, 0) + 1

        # 按频率排序并选择前N个
        sorted_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
        keywords = [word for word, freq in sorted_words[:max_keywords]]

        return keywords

    def get_sentences_needing_citations(
        self, result: DraftAnalysisResult, exclude_existing: bool = True
    ) -> List[Sentence]:
        """
        获取需要插入引用的句子

        Args:
            result: 草稿分析结果
            exclude_existing: 是否排除已有引用的句子

        Returns:
            需要引用的句子列表
        """
        if exclude_existing:
            return [s for s in result.sentences if not s.has_citation]
        else:
            return result.sentences

    def analyze_text_only(self, text: str) -> DraftAnalysisResult:
        """
        仅分析文本内容（非Word文档）

        Args:
            text: 纯文本内容

        Returns:
            草稿分析结果
        """
        result = DraftAnalysisResult()
        result.full_text = text
        result.paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]

        sentence_idx = 0
        for para_idx, paragraph in enumerate(result.paragraphs):
            sentences = self._split_sentences(paragraph)

            for sent_text in sentences:
                has_citation, citation_text = self._detect_citation(sent_text)
                keywords = self._extract_keywords(sent_text)

                sentence = Sentence(
                    text=sent_text,
                    index=sentence_idx,
                    paragraph_index=para_idx,
                    keywords=keywords,
                    has_citation=has_citation,
                    citation_text=citation_text,
                )

                result.sentences.append(sentence)
                sentence_idx += 1

        return result
