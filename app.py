import os
import sys
from pathlib import Path

# 添加src到路径
sys.path.insert(0, str(Path(__file__).parent / "src"))

import streamlit as st
import pandas as pd
from datetime import datetime

from src.literature.db_manager import (
    LiteratureDatabaseManager,
    create_literature_database,
)
from src.draft.analyzer import DraftAnalyzer
from src.citation.ai_matcher import (
    AICitationMatcher,
    AIAPIManager,
    SentenceWithAICitations,
)
from src.citation.format_learner import ReferenceFormatLearner
from src.utils.config import get_config


def init_session_state():
    """初始化session state"""
    if "db_manager" not in st.session_state:
        st.session_state.db_manager = None
    if "draft_analysis" not in st.session_state:
        st.session_state.draft_analysis = None
    if "citation_results" not in st.session_state:
        st.session_state.citation_results = None
    if "imported_files" not in st.session_state:
        st.session_state.imported_files = []


def render_sidebar():
    """渲染侧边栏"""
    with st.sidebar:
        st.title("📖 论文反插助手")
        st.caption("基于AI的学术论文引用自动插入工具")

        st.divider()

        # 1. API设置（可折叠）
        with st.expander("🔑 API配置", expanded=True):
            api_provider = st.selectbox(
                "API提供商",
                options=["deepseek", "openai", "anthropic"],
                index=0,
                help="选择AI模型提供商",
            )

            if api_provider == "deepseek":
                api_key = st.text_input(
                    "API密钥",
                    value="",
                    type="password",
                    placeholder="sk-...",
                    help="DeepSeek API密钥",
                )
                api_base_url = "https://api.deepseek.com/v1"
                model = st.selectbox(
                    "模型",
                    options=["deepseek-chat", "deepseek-reasoner"],
                    index=0,
                )
            elif api_provider == "openai":
                api_key = st.text_input(
                    "API密钥",
                    value="",
                    type="password",
                    placeholder="sk-...",
                )
                api_base_url = st.text_input(
                    "API地址（可选）",
                    value="",
                    placeholder="自定义中转地址",
                )
                model = st.selectbox(
                    "模型",
                    options=["gpt-4o", "gpt-4o-mini", "gpt-4-turbo"],
                    index=0,
                )
            else:
                api_key = st.text_input(
                    "API密钥",
                    value="",
                    type="password",
                    placeholder="sk-ant-...",
                )
                api_base_url = ""
                model = st.selectbox(
                    "模型",
                    options=["claude-3-5-sonnet-20241022"],
                    index=0,
                )

            # API状态检查
            if api_key:
                st.success("✅ API已配置")
            else:
                st.warning("⚠️ 请输入API密钥")

        # 2. 引用设置
        with st.expander("📚 引用设置"):
            citation_style = st.selectbox(
                "引用风格",
                options=["author-year", "numbered"],
                index=0,
                help="选择文中引用格式",
            )

            max_citations = st.slider(
                "每句最大引用数",
                min_value=1,
                max_value=5,
                value=2,
            )

            min_relevance = st.slider(
                "最低相关性阈值",
                min_value=0.0,
                max_value=1.0,
                value=0.6,
                step=0.05,
                help="低于此分数的引用将被忽略",
            )

            st.caption(f"当前阈值: {min_relevance:.2f} - 低于此分数的引用将被过滤")

        # 3. 检索引擎设置
        with st.expander("🔍 检索引擎"):
            use_hybrid_search = st.toggle(
                "启用混合检索",
                value=True,
                help="启用AI增强的混合检索（需要模型文件）",
            )

            if use_hybrid_search:
                st.markdown(
                    """
                <small style='color:green'>✅ 查询扩展 → 多路召回 → Cross-encoder重排 → MMR多样</small>
                """,
                    unsafe_allow_html=True,
                )
            else:
                st.markdown(
                    """
                <small style='color:orange'>⚠️ 仅使用关键词检索</small>
                """,
                    unsafe_allow_html=True,
                )

        # 4. 文献筛选策略
        with st.expander("⚖️ 文献筛选策略"):
            st.markdown("**两步筛选法**")
            st.caption("1. 语义筛选 → 2. 新颖度/引用加权排序")

            top_k_semantic = st.slider(
                "语义筛选保留数量",
                min_value=10,
                max_value=100,
                value=50,
                step=10,
            )

            st.divider()

            # 权重滑块
            col_w1, col_w2 = st.columns(2)
            with col_w1:
                weight_recency = st.slider(
                    "📅 新颖度",
                    0,
                    100,
                    50,
                    step=5,
                )
            with col_w2:
                weight_citation = 100 - weight_recency
                st.metric("📚 引用", f"{weight_citation}%")

            # 可视化权重
            st.progress(weight_recency / 100)
            st.caption(f"新颖度 {weight_recency}% | 引用 {weight_citation}%")

            # 预设按钮
            preset_col1, preset_col2 = st.columns(2)
            with preset_col1:
                if st.button("⚖️ 均衡", use_container_width=True):
                    weight_recency = 50
                    st.rerun()
            with preset_col2:
                if st.button("🆕 追新", use_container_width=True):
                    weight_recency = 80
                    st.rerun()

        # 5. 参考文献格式（可选）
        with st.expander("📝 参考文献格式", expanded=False):
            st.caption("粘贴目标期刊的参考文献示例，AI将学习格式")

            reference_example = st.text_area(
                "示例（可选）",
                height=100,
                placeholder="[1] Zhang, X. (2024). Title...",
            )

            if reference_example and st.button("🎓 学习格式", type="secondary"):
                if api_key:
                    with st.spinner("学习中..."):
                        temp_api_manager = AIAPIManager(
                            api_key=api_key,
                            base_url=api_base_url or "https://api.deepseek.com/v1",
                            model=model,
                            provider=api_provider,
                        )
                        format_learner = ReferenceFormatLearner(temp_api_manager)
                        learned_format = format_learner.learn_from_example(
                            reference_example
                        )
                        st.session_state.reference_format = learned_format
                        st.success(f"✅ 已学习: {learned_format.name}")
                else:
                    st.warning("⚠️ 请先配置API密钥")

        st.divider()

        # 数据库状态
        if st.session_state.db_manager:
            stats = st.session_state.db_manager.get_statistics()
            st.markdown(
                f"""
            <div style='padding:10px; background: #f0f2f6; border-radius:10px;'>
                <b>📊 数据库状态</b><br>
                文献数量: <b>{stats["total_papers"]}</b><br>
                最早文献: <b>{stats["earliest_year"]}</b><br>
                最新文献: <b>{stats["latest_year"]}</b>
            </div>
            """,
                unsafe_allow_html=True,
            )

        return {
            "api_provider": api_provider,
            "api_key": api_key,
            "api_base_url": api_base_url,
            "model": model,
            "citation_style": citation_style,
            "max_citations": max_citations,
            "min_relevance": min_relevance,
            "top_k_semantic": top_k_semantic,
            "weight_recency": weight_recency,
            "weight_citation": weight_citation,
            "reference_example": reference_example,
            "use_hybrid_search": use_hybrid_search,
        }


def render_literature_import():
    """渲染文献导入Tab"""
    st.markdown("### 📚 导入文献库")

    # 简洁的操作说明
    with st.expander("📋 操作指南", expanded=True):
        st.markdown("""
        **从Web of Science导入：**
        1. 在Web of Science中搜索文献 → 2. 选择要导出的文献
        3. 点击 **Export** → **Plain Text File** 
        4. 选择 **Full Record** 格式 → 5. 下载 .txt 文件
        6. 在下方上传文件
        """)

    # 文件上传区域
    st.markdown("**上传WOS导出文件**")
    uploaded_files = st.file_uploader(
        "拖拽文件到此处或点击选择",
        type=["txt"],
        accept_multiple_files=True,
        help="支持批量上传多个txt文件",
    )

    if uploaded_files:
        # 文件列表显示
        st.success(f"✅ 已选择 {len(uploaded_files)} 个文件")

        # 显示文件名
        with st.expander(f"查看文件列表 ({len(uploaded_files)}个)"):
            for f in uploaded_files:
                st.caption(f"📄 {f.name}")

        if st.button("🚀 开始导入", type="primary"):
            progress_bar = st.progress(0)
            status_text = st.empty()

            # 初始化数据库
            db_path = "data/literature.db"
            db_manager = LiteratureDatabaseManager(db_path)

            total_count = 0
            all_errors = []

            for idx, uploaded_file in enumerate(uploaded_files):
                status_text.text(f"正在导入: {uploaded_file.name}...")

                # 保存上传的文件
                file_path = f"uploads/{uploaded_file.name}"
                os.makedirs("uploads", exist_ok=True)
                with open(file_path, "wb") as f:
                    f.write(uploaded_file.getvalue())

                # 导入文献
                count, errors = db_manager.import_from_wos_txt(file_path)
                total_count += count
                all_errors.extend(errors)

                # 记录已导入的文件
                if uploaded_file.name not in st.session_state.imported_files:
                    st.session_state.imported_files.append(uploaded_file.name)

                progress_bar.progress((idx + 1) / len(uploaded_files))

            st.session_state.db_manager = db_manager

            # 显示统计
            stats = db_manager.get_statistics()

            st.success(f"✅ 成功导入 {total_count} 篇论文！")

            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("总文献数", stats["total_papers"])
            with col2:
                years = list(stats["year_distribution"].keys())
                if years:
                    st.metric("年份范围", f"{min(years)}-{max(years)}")
            with col3:
                st.metric("期刊种类", len(stats["top_journals"]))

            if all_errors:
                with st.expander(f"查看错误 ({len(all_errors)}个)"):
                    for error in all_errors[:10]:  # 只显示前10个
                        st.error(error)


def render_draft_upload():
    """渲染草稿上传Tab"""
    st.markdown("### 📝 上传草稿")

    # 检查是否已导入文献
    if st.session_state.db_manager is None:
        st.warning('⚠️ 请先在左侧"导入文献库"中导入文献')

        st.markdown("""
        ---
        **快速开始：**
        1. 切换到 **📚 导入文献库** 标签
        2. 上传Web of Science导出的.txt文件
        3. 等待导入完成
        """)
        return

    # 显示当前文献库信息
    stats = st.session_state.db_manager.get_statistics()
    st.success(f"✅ 已加载文献库: {stats['total_papers']} 篇论文")

    # 文件上传
    st.markdown("**上传Word文档**")
    uploaded_file = st.file_uploader(
        "拖拽文件到此处或点击选择",
        type=["docx"],
        help="上传写好但未插入引用的Word文档（.docx格式）",
    )

    if uploaded_file:
        # 文件信息
        st.info(f"📄 {uploaded_file.name} ({(uploaded_file.size / 1024):.1f} KB)")

        # 保存文件
        file_path = f"uploads/{uploaded_file.name}"
        os.makedirs("uploads", exist_ok=True)
        with open(file_path, "wb") as f:
            f.write(uploaded_file.getvalue())

        # 分析按钮
        if st.button("🔬 分析文档", type="primary"):
            with st.spinner("正在分析文档结构..."):
                analyzer = DraftAnalyzer()
                analysis = analyzer.analyze_draft(file_path)
                st.session_state.draft_analysis = analysis

            # 分析结果统计
            st.success(f"✅ 分析完成！")

            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("总句子数", len(analysis.sentences))
            with col2:
                needing_citations = len(
                    [s for s in analysis.sentences if not s.has_citation]
                )
                st.metric("需引用句子", needing_citations)
            with col3:
                st.metric("段落数", len(analysis.paragraphs))

            # 显示前几个句子
            with st.expander("预览句子"):
                for i, sent in enumerate(analysis.sentences[:5]):
                    st.markdown(f"**句子 {i + 1}:** {sent.text[:100]}...")
                    if sent.keywords:
                        st.caption(f"关键词: {', '.join(sent.keywords)}")


def render_citation_matching(config):
    """渲染引用匹配Tab"""
    st.header("🔍 引用匹配")

    # 检查前置条件
    if st.session_state.db_manager is None:
        st.warning("⚠️ 请先导入文献库")
        return

    if st.session_state.draft_analysis is None:
        st.warning("⚠️ 请先上传并分析草稿")
        return

    analysis = st.session_state.draft_analysis

    # 匹配选项
    col1, col2, col3 = st.columns(3)
    with col1:
        exclude_existing = st.checkbox(
            "跳过已有引用的句子", value=True, help="不处理已经包含引用的句子"
        )
    with col2:
        year_range = st.slider(
            "文献年份范围",
            min_value=5,
            max_value=30,
            value=10,
            help="只搜索最近N年的文献",
        )
    with col3:
        prioritize_recent = st.checkbox(
            "优先推荐新文献", value=True, help="优先匹配近5年的文献"
        )

    # 检查API配置
    if not config.get("api_key"):
        st.error("⚠️ 请在侧边栏输入API密钥")
        return

    # 开始匹配
    if st.button("开始AI匹配引用", type="primary"):
        sentences_to_match = analysis.sentences
        if exclude_existing:
            sentences_to_match = [s for s in analysis.sentences if not s.has_citation]

        if not sentences_to_match:
            st.warning("没有需要匹配的句子")
            return

        # 初始化AI API管理器
        api_manager = AIAPIManager(
            api_key=config["api_key"],
            base_url=config.get("api_base_url", "https://api.deepseek.com/v1"),
            model=config.get("model", "deepseek-chat"),
            provider=config.get("api_provider", "deepseek"),
        )

        # 初始化AI匹配器（传入用户设置的参数）
        matcher = AICitationMatcher(
            db_manager=st.session_state.db_manager,
            api_manager=api_manager,
            citation_style=config["citation_style"],
            max_citations=config["max_citations"],
            min_relevance=config.get("min_relevance", 0.6),
            batch_size=5,
            top_k_semantic=int(config.get("top_k_semantic", 50)),
            weight_recency=int(config.get("weight_recency", 50)),
            weight_citation=int(config.get("weight_citation", 50)),
            use_hybrid_search=config.get("use_hybrid_search", True),
        )

        st.info("🤖 正在使用AI进行语义匹配，这可能需要一些时间...")

        # 进度条
        progress_bar = st.progress(0)
        status_text = st.empty()

        def progress_callback(current, total):
            progress_bar.progress(current / total)
            status_text.text(f"正在AI匹配: 句子 {current}/{total}")

        # 批量匹配
        results = matcher.batch_match(
            sentences=sentences_to_match,
            year_range=year_range,
            progress_callback=progress_callback,
        )

        st.session_state.citation_results = results
        st.session_state.citation_matcher = matcher

        st.success(f"✅ AI匹配完成！共处理 {len(results)} 个句子")

        # 统计
        with_citations = len([r for r in results if r.citations])
        high_confidence = len(
            [r for r in results if any(c.confidence == "high" for c in r.citations)]
        )

        # 统计近5年文献占比
        current_year = datetime.now().year
        recent_papers = 0
        total_papers = 0
        for r in results:
            for c in r.citations:
                total_papers += 1
                if c.paper.year >= current_year - 5:
                    recent_papers += 1

        recent_ratio = (recent_papers / total_papers * 100) if total_papers > 0 else 0

        # 统计卡片（更美观的显示）
        st.markdown("### 📈 匹配统计")

        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("✅ 成功匹配", f"{with_citations}/{len(results)}")
        with col2:
            st.metric("🟢 高置信度", high_confidence)
        with col3:
            st.metric("📅 近5年文献", f"{recent_ratio:.0f}%")
        with col4:
            match_rate = with_citations / len(results) * 100 if len(results) > 0 else 0
            st.metric("📊 匹配率", f"{match_rate:.1f}%")

        # 进度条可视化
        if total_papers > 0:
            st.progress(match_rate / 100)
            st.caption(f"已为 {with_citations} 个句子找到合适的引用")


def render_results_review(config):
    """渲染结果查看Tab"""
    st.header("📊 查看与导出")

    if st.session_state.citation_results is None:
        st.warning("⚠️ 请先完成引用匹配")
        return

    results = st.session_state.citation_results
    matcher = st.session_state.get("citation_matcher")

    # 导出选项
    st.subheader("导出设置")
    col1, col2 = st.columns(2)
    with col1:
        output_format = st.selectbox(
            "输出格式", options=["Word文档", "Markdown", "纯文本"], index=0
        )
    with col2:
        bibliography_style = st.selectbox(
            "参考文献格式", options=["apa", "nature", "vancouver", "ieee"], index=0
        )

    # 显示当前筛选策略配置
    with st.expander("📊 当前筛选策略", expanded=True):
        top_k = config.get("top_k_semantic", 50)
        w_rec = config.get("weight_recency", 50)
        w_cit = config.get("weight_citation", 50)

        st.markdown(f"**第一步：语义筛选** - 选出最相关的前 **{top_k}** 篇")

        st.markdown("**第二步：加权排序**")
        col_w1, col_w2 = st.columns(2)
        with col_w1:
            st.metric("📅 新颖度权重", f"{w_rec}%")
        with col_w2:
            st.metric("📚 引用次数权重", f"{w_cit}%")

        if w_rec + w_cit == 100:
            st.success("✅ 权重分配正确")

    st.divider()

    # 显示匹配结果
    st.subheader("匹配结果详情")

    # 过滤选项
    show_only_with_citations = st.checkbox("只显示有引用的句子", value=False)

    display_results = results
    if show_only_with_citations:
        display_results = [r for r in results if r.citations]

    # 分页显示
    page_size = 10
    total_pages = max(1, (len(display_results) + page_size - 1) // page_size)
    page = (
        st.number_input(
            f"页码 (共{total_pages}页)", min_value=1, max_value=total_pages, value=1
        )
        - 1
    )

    start_idx = page * page_size
    end_idx = min(start_idx + page_size, len(display_results))

    for i, result in enumerate(display_results[start_idx:end_idx], start=start_idx + 1):
        with st.container():
            st.markdown(f"**句子 {i}**")
            st.info(result.sentence.text)

            if result.sentence.has_citation:
                st.success(f"✓ 已有引用: {result.sentence.citation_text}")
            elif result.citations:
                st.markdown("**AI推荐引用:**")
                for j, citation in enumerate(result.citations, 1):
                    paper = citation.paper
                    col1, col2 = st.columns([3, 1])

                    # 计算年份标签
                    current_year_now = datetime.now().year
                    year_diff = current_year_now - paper.year
                    if year_diff <= 2:
                        year_badge = "🔥 最新"
                    elif year_diff <= 5:
                        year_badge = "⭐ 近5年"
                    elif year_diff <= 10:
                        year_badge = "📚 近10年"
                    else:
                        year_badge = "📖 经典"

                    with col1:
                        st.markdown(f"{j}. **{paper.title}**")
                        st.caption(f"作者: {paper.authors[:100]}...")
                        st.caption(
                            f"期刊: {paper.journal} | {year_badge} ({paper.year}) | 被引: {paper.cited_by}次"
                        )
                        confidence_emoji = {"high": "🟢", "medium": "🟡", "low": "🔴"}
                        emoji = confidence_emoji.get(citation.confidence, "⚪")

                        # AI评分和置信度
                        score_color = (
                            "green"
                            if citation.relevance_score >= 0.75
                            else "orange"
                            if citation.relevance_score >= 0.5
                            else "red"
                        )
                        st.markdown(
                            f"<span style='color:{score_color}'>{emoji} AI评分: {citation.relevance_score:.2f}</span> "
                            f"<span style='color:gray'>(置信度: {citation.confidence})</span>",
                            unsafe_allow_html=True,
                        )

                        # 匹配理由 - 更详细的显示
                        if citation.relevance_reason:
                            with st.expander("📝 查看匹配理由", expanded=False):
                                st.markdown(f"_{citation.relevance_reason}_")
                    with col2:
                        cite_text = (
                            matcher.format_citation(citation, j)
                            if matcher
                            else f"[{j}]"
                        )
                        st.code(cite_text)
            else:
                st.warning("AI未找到相关文献")

            st.divider()

    # 参考文献序号格式设置
    st.subheader("参考文献序号格式")
    ref_numbering = st.radio(
        "选择序号格式",
        options=["numbered", "none", "author_year"],
        format_func=lambda x: {
            "numbered": "[1], [2], [3]...",
            "none": "无序号（直接列出）",
            "author_year": "(Author, Year)",
        }.get(x, x),
        help="选择参考文献列表的编号方式",
    )

    # 导出按钮
    st.subheader("导出文档")

    if st.button("生成带引用的文档", type="primary"):
        with st.spinner("正在生成文档..."):
            # 保持段落结构重建文档
            from src.draft.analyzer import DraftAnalyzer

            # 构建段落映射（按段落索引组织句子）
            paragraph_map = {}
            for result in results:
                para_idx = result.sentence.paragraph_index
                if para_idx not in paragraph_map:
                    paragraph_map[para_idx] = []
                paragraph_map[para_idx].append(result)

            # 按段落重建文本
            paragraphs_text = []
            for para_idx in sorted(paragraph_map.keys()):
                para_sentences = paragraph_map[para_idx]
                para_text_parts = []

                for result in para_sentences:
                    if (
                        result.citations
                        and not result.sentence.has_citation
                        and matcher
                    ):
                        # 插入引用
                        new_text = matcher.insert_citations_into_text(
                            result.sentence, result.citations
                        )
                        para_text_parts.append(new_text)
                    else:
                        para_text_parts.append(result.sentence.text)

                # 组合成段落（保留原段落结构）
                paragraph_text = " ".join(para_text_parts)
                paragraphs_text.append(paragraph_text)

            # 用段落分隔符连接
            full_text = "\n\n".join(paragraphs_text)

            # 添加参考文献
            if matcher:
                # 检查是否有学习的格式
                learned_format = st.session_state.get("reference_format")
                if learned_format and config.get("api_key"):
                    # 使用学习的格式
                    with st.spinner("正在使用学习到的格式生成参考文献..."):
                        api_manager = AIAPIManager(
                            api_key=config["api_key"],
                            base_url=config.get(
                                "api_base_url", "https://api.deepseek.com/v1"
                            ),
                            model=config.get("model", "deepseek-chat"),
                            provider=config.get("api_provider", "deepseek"),
                        )
                        format_learner = ReferenceFormatLearner(api_manager)
                        format_learner.format_cache = learned_format

                        # 收集所有使用过的论文
                        used_papers = {}
                        for swc in results:
                            for citation in swc.citations:
                                paper_id = citation.paper.id
                                if paper_id not in used_papers:
                                    used_papers[paper_id] = citation.paper

                        # 使用学习的格式批量格式化
                        sorted_papers = sorted(
                            used_papers.values(),
                            key=lambda p: (
                                p.authors.split(",")[0].strip().split()[-1]
                                if p.authors
                                else ""
                            ).lower(),
                        )

                        formatted_refs = format_learner.batch_format(sorted_papers)

                        # 根据序号格式生成参考文献
                        if ref_numbering == "numbered":
                            bibliography = "# References\n\n" + "\n\n".join(
                                f"[{i + 1}] {ref}"
                                for i, ref in enumerate(formatted_refs)
                            )
                        elif ref_numbering == "author_year":
                            bibliography = "# References\n\n" + "\n\n".join(
                                formatted_refs
                            )
                        else:  # none
                            bibliography = "# References\n\n" + "\n\n".join(
                                formatted_refs
                            )
                else:
                    # 使用默认格式
                    used_papers = {}
                    for swc in results:
                        for citation in swc.citations:
                            paper_id = citation.paper.id
                            if paper_id not in used_papers:
                                used_papers[paper_id] = citation.paper

                    if used_papers:
                        sorted_papers = sorted(
                            used_papers.values(),
                            key=lambda p: (
                                p.authors.split(",")[0].strip().split()[-1]
                                if p.authors
                                else ""
                            ).lower(),
                        )

                        # 根据序号格式生成参考文献
                        if ref_numbering == "numbered":
                            bibliography = "# References\n\n"
                            for i, paper in enumerate(sorted_papers, 1):
                                authors = paper.authors.replace(";", ", ")
                                ref = f"[{i}] {authors} ({paper.year}). {paper.title}. {paper.journal}, {paper.volume}({paper.issue}), {paper.pages}."
                                bibliography += ref + "\n\n"
                        elif ref_numbering == "author_year":
                            bibliography = "# References\n\n"
                            for paper in sorted_papers:
                                authors = paper.authors.replace(";", ", ")
                                ref = f"{authors} ({paper.year}). {paper.title}. {paper.journal}, {paper.volume}({paper.issue}), {paper.pages}."
                                bibliography += ref + "\n\n"
                        else:  # none
                            bibliography = "# References\n\n"
                            for paper in sorted_papers:
                                authors = paper.authors.replace(";", ", ")
                                ref = f"{authors} ({paper.year}). {paper.title}. {paper.journal}, {paper.volume}({paper.issue}), {paper.pages}."
                                bibliography += ref + "\n\n"
                    else:
                        bibliography = "# References\n\n暂无引用文献"

                full_text += "\n\n" + bibliography.strip()

            # 确保output目录存在
            os.makedirs("output", exist_ok=True)

            # 保存文件
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

            if output_format == "纯文本":
                output_path = f"output/cited_draft_{timestamp}.txt"
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(full_text)

                with open(output_path, "r", encoding="utf-8") as f:
                    st.download_button(
                        label="下载文本文件",
                        data=f.read(),
                        file_name=f"cited_draft_{timestamp}.txt",
                        mime="text/plain",
                    )

            elif output_format == "Markdown":
                output_path = f"output/cited_draft_{timestamp}.md"
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(full_text)

                with open(output_path, "r", encoding="utf-8") as f:
                    st.download_button(
                        label="下载Markdown文件",
                        data=f.read(),
                        file_name=f"cited_draft_{timestamp}.md",
                        mime="text/markdown",
                    )

            else:  # Word文档
                from docx import Document
                from docx.shared import Pt
                from docx.oxml.ns import qn

                output_path = f"output/cited_draft_{timestamp}.docx"
                doc = Document()

                def set_times_new_roman(run):
                    """设置Times New Roman字体"""
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    run.font.size = Pt(12)

                # 添加内容（保持段落结构）
                for para_idx in sorted(paragraph_map.keys()):
                    para_sentences = paragraph_map[para_idx]
                    para_text_parts = []

                    for result in para_sentences:
                        if (
                            result.citations
                            and not result.sentence.has_citation
                            and matcher
                        ):
                            new_text = matcher.insert_citations_into_text(
                                result.sentence, result.citations
                            )
                            para_text_parts.append(new_text)
                        else:
                            para_text_parts.append(result.sentence.text)

                    # 添加段落
                    paragraph_text = " ".join(para_text_parts)
                    p = doc.add_paragraph(paragraph_text)

                    # 设置字体
                    for run in p.runs:
                        set_times_new_roman(run)

                # 添加参考文献
                if matcher:
                    doc.add_heading("References", level=1)

                    # 收集所有使用过的论文
                    used_papers = {}
                    for swc in results:
                        for citation in swc.citations:
                            paper_id = citation.paper.id
                            if paper_id not in used_papers:
                                used_papers[paper_id] = citation.paper

                    if used_papers:
                        sorted_papers = sorted(
                            used_papers.values(),
                            key=lambda p: (
                                p.authors.split(",")[0].strip().split()[-1]
                                if p.authors
                                else ""
                            ).lower(),
                        )

                        # 检查是否有学习的格式
                        learned_format = st.session_state.get("reference_format")
                        if learned_format and config.get("api_key"):
                            # 使用学习的格式
                            api_manager = AIAPIManager(
                                api_key=config["api_key"],
                                base_url=config.get(
                                    "api_base_url", "https://api.deepseek.com/v1"
                                ),
                                model=config.get("model", "deepseek-chat"),
                                provider=config.get("api_provider", "deepseek"),
                            )
                            format_learner = ReferenceFormatLearner(api_manager)
                            format_learner.format_cache = learned_format
                            formatted_refs = format_learner.batch_format(sorted_papers)
                        else:
                            # 使用默认格式
                            formatted_refs = []
                            for paper in sorted_papers:
                                authors = paper.authors.replace(";", ", ")
                                ref = f"{authors} ({paper.year}). {paper.title}. {paper.journal}, {paper.volume}({paper.issue}), {paper.pages}."
                                formatted_refs.append(ref)

                        # 根据序号格式添加参考文献
                        for i, ref in enumerate(formatted_refs, 1):
                            if ref_numbering == "numbered":
                                p = doc.add_paragraph(f"[{i}] {ref}")
                            elif ref_numbering == "author_year":
                                # 从引用中提取作者-年份格式
                                p = doc.add_paragraph(ref)
                            else:  # none
                                p = doc.add_paragraph(ref)

                            # 设置字体
                            for run in p.runs:
                                set_times_new_roman(run)

                doc.save(output_path)

                with open(output_path, "rb") as f:
                    st.download_button(
                        label="下载Word文档",
                        data=f.read(),
                        file_name=f"cited_draft_{timestamp}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )

            st.success(f"✅ 文档已生成: {output_path}")


def main():
    """主函数"""
    st.set_page_config(
        page_title="论文反插助手",
        page_icon="📚",
        layout="wide",
        initial_sidebar_state="expanded",
    )

    # 页面样式
    st.markdown(
        """
    <style>
    .main-header {
        font-size: 2rem;
        font-weight: bold;
        margin-bottom: 1rem;
    }
    .sub-header {
        font-size: 1.2rem;
        font-weight: 600;
        margin-top: 1rem;
    }
    .highlight-box {
        padding: 1rem;
        border-radius: 0.5rem;
        background: #f0f2f6;
        margin: 0.5rem 0;
    }
    </style>
    """,
        unsafe_allow_html=True,
    )

    # 页面标题
    st.markdown('<p class="main-header">📚 论文反插助手</p>', unsafe_allow_html=True)
    st.markdown("基于AI的智能引用插入工具 | 上传文献库和草稿，自动匹配最相关的参考文献")

    # 初始化session state
    init_session_state()

    # 渲染侧边栏
    config = render_sidebar()

    # 创建标签页（带图标）
    tab1, tab2, tab3, tab4 = st.tabs(
        ["📚 导入文献库", "📝 上传草稿", "⚡ AI匹配", "📊 查看与导出"]
    )

    with tab1:
        render_literature_import()

    with tab2:
        render_draft_upload()

    with tab3:
        render_citation_matching(config)

    with tab4:
        render_results_review(config)

    # 页脚
    st.sidebar.markdown("---")
    st.sidebar.caption("论文反插助手 v1.0")


if __name__ == "__main__":
    main()
